import os
import pdfplumber
from docx import Document
import openpyxl
import pandas as pd
from PIL import Image
import pytesseract

def detect_and_extract_text(file_path: str, file_extension: str) -> dict:
    """
    Detects the document format and automatically routes it to the correct 
    processing path to extract structured text or tabular elements.
    """
    ext = file_extension.lower().replace(".", "")
    
    if ext in ["pdf"]:
        return process_pdf(file_path)
    elif ext in ["docx"]:
        return process_docx(file_path)
    elif ext in ["xlsx"]:
        return process_xlsx(file_path)
    elif ext in ["jpg", "jpeg", "png"]:
        return process_image(file_path)
    else:
        raise ValueError(f"Unsupported file format: .{ext}")

def process_pdf(file_path: str) -> dict:
    extracted_data = {"type": "digital_pdf", "pages": []}
    
    with pdfplumber.open(file_path) as pdf:
        # Simple heuristic to check if the PDF is scanned or digital
        first_page_text = pdf.pages[0].extract_text() or ""
        if len(first_page_text.strip()) < 50:
            # Re-route to OCR processing if there's little to no digital text layer
            return process_scanned_pdf(file_path)
            
        for index, page in enumerate(pdf.pages):
            page_data = {
                "page_number": index + 1,
                "text": page.extract_text(),
                "tables": page.extract_tables()  # Preserves structure
            }
            extracted_data["pages"].append(page_data)
            
    return extracted_data

def process_docx(file_path: str) -> dict:
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Simple structured extraction
    tables_data = []
    for table in doc.tables:
        current_table = []
        for row in table.rows:
            current_table.append([cell.text.strip() for cell in row.cells])
        tables_data.append(current_table)
        
    return {
        "type": "docx",
        "content": paragraphs,
        "tables": tables_data
    }

def process_xlsx(file_path: str) -> dict:
    excel_file = pd.ExcelFile(file_path)
    valid_sheets = {}
    
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        
        # Heuristic: If a sheet has few rows or mostly empty columns, flag as metadata/chart sheets
        if df.empty or df.shape[0] < 2 or df.shape[1] < 2:
            continue  # Ignore metadata or raw charts
            
        # Standardize structured data into JSON records
        valid_sheets[sheet_name] = df.dropna(how='all').to_dict(orient="records")
        
    return {
        "type": "xlsx",
        "sheets": valid_sheets
    }

def process_image(file_path: str) -> dict:
    image = Image.open(file_path)
    
    # Retrieve detailed data including word-level confidence metrics
    ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    
    # Extract confidence values ignoring structural space characters (-1)
    confidences = [int(c) for c in ocr_data['conf'] if c != -1]
    mean_confidence = sum(confidences) / len(confidences) if confidences else 0
    
    # Extract compiled string output
    raw_text = pytesseract.image_to_string(image)
    
    # Post-process common artefacts (line breaks splitting words)
    clean_text = raw_text.replace("-\n", "").replace("\n", " ")
    
    return {
        "type": "scanned_image",
        "text": clean_text,
        "ocr_confidence": mean_confidence,
        "is_low_quality": mean_confidence < 65.0  # Configurable threshold flag
    }

def process_scanned_pdf(file_path: str) -> dict:
    # In a full run, convert PDF pages to images via pdf2image, then process via process_image
    # For now, stub as a fallback structure
    return {"type": "scanned_pdf", "text": "Scanned pipeline routing triggered."}